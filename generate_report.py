# -*- coding: utf-8 -*-
"""生成流光项目成果报告 Word 文档（聚焦白皮书第二、三章）"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

def set_cn(run):
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return run

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_cn(r)
    return h

def add_p(text, bold=False):
    p = doc.add_paragraph()
    r = set_cn(p.add_run(text))
    r.font.size = Pt(11)
    if bold:
        r.font.bold = True
    return p

def add_b(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = set_cn(p.add_run(text))
    r.font.size = Pt(11)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = set_cn(c.paragraphs[0].add_run(h))
        r.font.bold = True
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = set_cn(c.paragraphs[0].add_run(str(val)))
            r.font.size = Pt(10)
    doc.add_paragraph()
    return t

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = set_cn(p.add_run('流光智能数据资产平台'))
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = set_cn(p.add_run('数据处理全流程成果报告'))
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = set_cn(p.add_run('对齐白皮书章节：第二章（数据处理）、第三章（字段存储与分类）'))
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = set_cn(p.add_run('编制人：梓旬\n2026年3月2日'))
r.font.size = Pt(12)

doc.add_page_break()

# ===== 一、我负责的部分说明 =====
add_h('一、对齐说明', 1)
add_p('本报告对齐《技术白皮书（数据处理全流程）》中的以下两个章节，结合流光平台的实际开发成果进行说明：')

add_table(
    ['白皮书章节', '流光平台对应模块', '实现状态'],
    [
        ['二、数据处理：精准提取 + 溯源 + 适配多格式', 'LLM Schema 映射 + Embedding 向量化 + ETL Transformer', '已完成'],
        ['三、字段存储与分类：分类清晰 + 权限可控 + 动态更新', 'PostgreSQL 分层存储 + RBAC 权限 + 知识图谱关联', '已完成'],
    ]
)

doc.add_page_break()

# ===== 二、数据处理 =====
add_h('二、数据处理（对齐白皮书第二章）', 1)
add_p('白皮书定义："提取得准、来源可溯、格式适配"，针对不同类型数据采用差异化方案，同时保障数据原始性。')

# --- 2.1 ---
add_h('2.1 智能字段提取与溯源', 2)
add_p('白皮书要求：通过自然语言理解模型识别并提取业务所需的核心字段，提取后自动标记字段溯源信息。')
add_p('流光实现：', bold=True)
add_b('使用 LLM（大语言模型）驱动 Schema 映射引擎，将飞书多维表格的原始字段自动映射为平台标准数据结构（如标题、内容、分类等）')
add_b('映射结果保留字段来源信息，可追溯到原始表格的具体列名和记录 ID（feishu_record_id）')
add_b('引入 MD5 哈希缓存机制：对相同 Schema 结构不重复调用 LLM，相同表结构只映射一次，后续直接复用缓存结果')
add_b('无法映射的字段不丢弃，自动存入 asset_tags（JSONB 字段），保留原始数据完整性')

# --- 2.2 ---
add_h('2.2 标准结构化数据提取（表格类）', 2)
add_p('白皮书要求：针对 Excel、CSV 等表格类数据，直接使用工具读取，全程不经过 AI 模型加工，确保数据与原始完全一致。')
add_p('流光实现：', bold=True)
add_b('飞书多维表格本身就是结构化数据，通过飞书 API 直接读取原始字段值')
add_b('标准字段（标题、内容、创建时间等）直接映射入库，不经过 LLM 修改，保证数据原始性')
add_b('ETL Transformer 模块中对标准字段做直通处理，仅对需要智能识别的字段才调用 LLM')

# --- 2.3 ---
add_h('2.3 非结构化数据处理（PDF / 图片）', 2)
add_p('白皮书要求：自动识别 PDF 类型，对纯文本 PDF 直接提取，对扫描版启用 OCR；图片统一通过高精度 OCR 识别。')
add_p('流光实现：', bold=True)
add_b('平台支持文件上传处理（file_upload 服务），可接收多种格式文件')
add_b('云文档导入服务（cloud_doc_import）支持从飞书云文档中提取文本内容')
add_b('当前主数据源为飞书多维表格（结构化为主），PDF/OCR 能力作为扩展预留')

# --- 2.4 ---
add_h('2.4 数据去重与降噪', 2)
add_p('白皮书要求：通过哈希算法对比数据指纹，自动剔除重复数据，过滤无效噪声。')
add_p('流光实现：', bold=True)
add_b('Upsert 机制：以 feishu_record_id 为唯一键，同一条数据重复入库时自动更新而非重复插入，从根源杜绝重复')
add_b('Schema 映射缓存基于 MD5 哈希比对，相同结构不重复处理')
add_b('ETL Loader 在入库前自动过滤空值、缺失关键字段的无效数据')

# --- 2.5 Embedding ---
add_h('2.5 Embedding 向量化（白皮书之外的增值实现）', 2)
add_p('在白皮书数据处理基础上，流光平台额外实现了文本向量化能力，为后续智能检索提供基础：')
add_b('调用 LLM Embedding 模型将文本内容转为 1536 维语义向量')
add_b('向量存储于 PostgreSQL pgvector 扩展，支持余弦相似度检索')
add_b('批量处理时自动控制 API 调用频率，防止限流')

# --- 2.6 Skill & MCP ---
add_h('2.6 核心技能封装与 MCP 接口', 2)

add_p('Skill Package（核心技能封装）：', bold=True)
add_p('当前数据处理相关能力（字段提取、结构化解析、非结构化处理）已在 ETL 服务层以模块化方式实现。下一阶段计划将这些能力抽象封装为独立的 Skill Package，支持按需调用和灵活编排。')

add_p('MCP 接口封装：', bold=True)
add_p('白皮书定义了数据处理相关的 3 个 MCP 接口，流光平台对应实现情况如下：')

add_table(
    ['白皮书 MCP 接口', '接口功能', '流光对应实现', '状态'],
    [
        ['接口1：核心字段提取接口',
         '接收原始数据及目标字段规则，自动完成字段提取、溯源标记',
         'ETL Transformer 模块内置，LLM 自动完成 Schema 映射与字段提取，返回带溯源的结果',
         '已实现'],
        ['接口2：结构化数据解析接口',
         '针对 Excel/CSV 等表格类文件精准解析',
         '飞书 API 直读 + Transformer 直通映射，结构化数据不经 LLM 加工，保证原始性',
         '已实现'],
        ['接口3：非结构化数据处理接口',
         '统一处理 PDF、图片，自动判断类型并执行提取',
         'POST /api/upload 文件上传接口 + cloud_doc_import 云文档导入服务',
         '基础已实现'],
    ]
)

doc.add_page_break()

# ===== 三、字段存储与分类 =====
add_h('三、字段存储与分类（对齐白皮书第三章）', 1)
add_p('白皮书定义："分类存储、权限管理、关联更新"，让数据既规整又能灵活复用。')

# --- 3.1 ---
add_h('3.1 字段分类与标签', 2)
add_p('白皮书要求：根据业务场景预设分类规则，提取后的字段自动归入对应类别，同时为每个字段打上业务标签。')
add_p('流光实现：', bold=True)
add_b('三大核心资产类型：Document（文档）、Meeting（会议）、ChatMessage（聊天消息），ETL 抽取时按数据源自动归类')
add_b('JSONB 标签字段 asset_tags：LLM 无法标准映射的扩展字段自动存入标签，支持灵活的业务分类检索')
add_b('多维资产扩展：在三大核心类型之外，还支持待办（Todo）、报告（Report）、知识图谱节点（KnowledgeGraph）等扩展类型')

# --- 3.2 ---
add_h('3.2 分层存储与权限管控', 2)
add_p('白皮书要求：初期采用 PostgreSQL 存储个人/测试数据；线上部署后按角色设置数据调用权限，避免泄露或误操作。')
add_p('流光实现：', bold=True)
add_b('数据库选型 PostgreSQL 16 + pgvector 扩展，兼顾关系型存储与向量检索')
add_b('三级 RBAC 角色体系：employee（普通员工）、executive（管理层）、admin（系统管理员）')
add_b('行级安全（RLS）：所有数据查询强制附加 owner_id 过滤条件——普通员工只能看到自己的数据，管理员可查看全部')
add_b('RAG 智能问答也严格遵守权限隔离，回答仅基于当前用户有权访问的数据')

add_table(
    ['角色', '数据可见范围', '操作权限'],
    [
        ['employee', '仅个人数据', '查看、智能问答'],
        ['executive', '部门及个人数据', '查看、分析、报告'],
        ['admin', '全局数据', '全部操作 + ETL 管理 + 用户管理'],
    ]
)

# --- 3.3 ---
add_h('3.3 数据关联与动态更新', 2)
add_p('白皮书要求：建立字段间的关联关系，新数据入库时系统自动匹配相同类型的已有数据，按预设规则更新字段。')
add_p('流光实现：', bold=True)
add_b('知识图谱构建服务（kg_builder）：自动发现数据实体间的关联关系，构建可视化的知识网络')
add_b('Upsert 更新策略：同一条飞书记录（feishu_record_id）再次入库时自动覆盖旧数据，保证数据时效性')
add_b('增量同步：每次 ETL 只拉取上次同步之后变更的数据，避免全量覆盖，旧数据安全保留')

# --- 3.4 ---
add_h('3.4 数据检索', 2)
add_p('白皮书要求：支持精准检索与模糊检索。')
add_p('流光实现：', bold=True)
add_b('向量语义检索：pgvector 余弦相似度搜索，用户用自然语言提问就能找到相关数据')
add_b('关键词精准检索：PostgreSQL 全文搜索（BM25 算法），支持关键词精确匹配')
add_b('RRF 融合排序：Reciprocal Rank Fusion 算法把两路检索结果合并排序，兼顾语义理解和关键词匹配')
add_b('RAG 智能问答：检索到相关数据后，LLM 生成自然语言回答，通过 SSE 流式推送给用户')

# --- 3.5 Skill & MCP ---
add_h('3.5 核心技能封装与 MCP 接口', 2)

add_p('Skill Package（核心技能封装）：', bold=True)
add_p('当前字段存储、权限管控、关联更新、数据检索等能力已在服务层实现。下一阶段计划将分类存储、权限配置、检索等核心能力封装为标准化 Skill Package，便于跨项目复用和统一调度。')

add_p('MCP 接口封装：', bold=True)
add_p('白皮书定义了字段存储与分类相关的 4 个 MCP 接口，流光平台对应实现情况如下：')

add_table(
    ['白皮书 MCP 接口', '接口功能', '流光对应实现', '状态'],
    [
        ['接口1：字段分类存储接口',
         '接收带分类标签的字段数据，完成分层存储',
         'ETL Loader 模块，自动按资产类型分类入库，扩展字段存入 asset_tags',
         '已实现'],
        ['接口2：权限配置接口',
         '为不同角色配置数据访问、操作权限',
         'PATCH /api/users/{id}/role 角色管理接口 + deps.py 权限守卫中间件',
         '已实现'],
        ['接口3：数据关联更新接口',
         '新数据入库时自动匹配已有数据并更新',
         'ETL Upsert 自动覆盖更新 + 知识图谱关联构建',
         '已实现'],
        ['接口4：数据检索接口',
         '根据标签、关键词等检索目标数据',
         'POST /api/chat/stream 混合检索 + RAG 问答；GET /api/assets/list 列表检索',
         '已实现'],
    ]
)

doc.add_page_break()

# ===== 四、总结 =====
add_h('四、个人成果总结', 1)

add_p('本人在流光平台项目中，主要负责白皮书第二章（数据处理）和第三章（字段存储与分类）对应的功能实现，核心成果如下：')

add_p('数据处理方面（对齐白皮书第二章）：', bold=True)
add_b('完成了 LLM 驱动的智能字段提取与 Schema 映射引擎，支持自动识别和提取业务字段')
add_b('实现了 MD5 缓存机制，避免重复调用 LLM，提升效率')
add_b('完成了结构化数据直通处理，保证数据原始性不被模型篡改')
add_b('实现了 Upsert 去重机制和数据降噪处理')
add_b('完成了 Embedding 向量化能力，为智能检索打下基础')

add_p('字段存储与分类方面（对齐白皮书第三章）：', bold=True)
add_b('完成了基于 PostgreSQL 的分层存储方案和 JSONB 标签分类体系')
add_b('实现了三级 RBAC 权限模型（employee / executive / admin）和行级安全（RLS）')
add_b('完成了知识图谱关联构建和 Upsert 动态更新机制')
add_b('实现了向量 + BM25 + RRF 的混合检索引擎，支持精准和模糊两种检索模式')

add_p('')
add_p('白皮书对齐覆盖度汇总：', bold=True)
add_table(
    ['白皮书要求项', '实现情况'],
    [
        ['智能字段提取与溯源', '已完成，LLM Schema 映射 + 溯源标记'],
        ['结构化数据原始性保障', '已完成，标准字段直通不经 LLM'],
        ['非结构化数据处理', '基础完成，文件上传 + 云文档导入'],
        ['数据去重与降噪', '已完成，Upsert + MD5 缓存'],
        ['字段分类与标签', '已完成，三大资产类型 + JSONB 标签'],
        ['分层存储与权限管控', '已完成，PostgreSQL + 三级 RBAC + RLS'],
        ['数据关联与动态更新', '已完成，知识图谱 + Upsert 覆盖更新'],
        ['精准与模糊检索', '已完成，向量 + BM25 + RRF 混合检索'],
        ['第二章 MCP 接口（3个）', '已实现'],
        ['第三章 MCP 接口（4个）', '已实现'],
        ['Skill Package 技能封装', '下一阶段计划封装'],
    ]
)

# ===== 保存 =====
output_path = r'D:\CC\liuguang-data-platform\梓旬_数据处理与存储成果报告.docx'
doc.save(output_path)
print(f'Done: {output_path}')
